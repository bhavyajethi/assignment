import os
from typing import Dict, Any, Optional 
import PyPDF2
import pdfplumber
from docx import Document
from ..services.ai_service import AIService

class ResumeParser:
    """Parses different resume file formats and uses LLM for content extraction."""
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', 'doc']
        self.ai_service = AIService()

    async def parse_resume(self, file_path:str) -> Dict[str, Any]:
        """Parse the resume file from file path and extract structured data using LLM."""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File Not Found: {file_path}")

            _,ext = os.path.splitext(file_path.lower())
            if ext == ".pdf":
                raw_text = self._extract_from_pdf(file_path)
            elif ext in ["docx", "doc"]:
                raw_text = self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")

            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError("Resume is too short or empty.")

            # Use llm to extract the structured data
            structured_data = await self.ai_service.parse_resume_content(raw_text)
            structured_data['text'] = raw_text
            structured_data['file_type'] = ext
            structured_data['word_length'] = len(raw_text.split())
            return structured_data

        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}")

    def _extract_from_pdf(self, file_path:str) -> str:
        """extract raw text from pdf"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text+=page_text + "\n"

            "if pdfplumber fails, fallback to PyPDF2"
            if len(text.strip()) < 50:
                text = ""
                with open('file_path', 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text+=page_text + "\n"      

        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
        return text.strip()

    def _extract_from_docx(self, file_path:str) -> str:
        """extract raw text from DOCX/DOC"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text.strip()      

        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")